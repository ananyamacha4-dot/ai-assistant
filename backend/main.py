
import smtplib

import asyncio

from email.mime.text import MIMEText

from email.mime.multipart import MIMEMultipart
import io
import sys
import textwrap
from fastapi import (
     FastAPI,
    UploadFile,
    File
)
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.tools import tool

from groq import Groq

from database import SessionLocal, engine
from models import User, Base
from auth import (
    hash_password,
    verify_password,
    create_token
)

import logging
import os
import shutil

from pypdf import PdfReader

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ai-chatbot")

# =========================
# LOAD ENV
# =========================

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY   = os.getenv("GROQ_API_KEY")
PDF_CONTEXT_FILE = "pdf_context.txt"

# Lazy Groq client — only built if the key exists.
def get_groq_client():

    if not GROQ_API_KEY:

        return None

    return Groq(api_key=GROQ_API_KEY)

# =========================
# CREATE DATABASE TABLES
# =========================

Base.metadata.create_all(bind=engine)

# =========================
# FASTAPI
# =========================

app = FastAPI()

# =========================
# CORS
# =========================

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "https://ai-assistant-gwco.vercel.app",
        "https://ai-assistant-gwco-cxdzrjncv-ananyamacha4-dots-projects.vercel.app",
    ],
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# LIGHTWEIGHT AI HELPERS
# =========================

# Ordered fallback chain — tried in order until one succeeds.
# gemini-2.5-flash-lite has the highest free-tier RPD (1000/day),
# so we start there and fall back to other variants on quota errors.
GEMINI_MODEL_CHAIN = [
    "gemini-2.5-flash-lite",
    "gemini-flash-latest",
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
]

_gemini_configured = False

def _ensure_gemini_configured():
    global _gemini_configured
    if not _gemini_configured and GEMINI_API_KEY:
        genai.configure(api_key=GEMINI_API_KEY)
        _gemini_configured = True

def get_gemini_model(model_name: str = "gemini-2.5-flash-lite"):
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set in environment variables.")
    _ensure_gemini_configured()
    return genai.GenerativeModel(model_name)

def _extract_response_text(response) -> str:

    try:
        text = response.text or ""
        if text.strip():
            return text
    except Exception as text_err:
        logger.warning("response.text raised: %s", text_err)

    try:
        candidates = getattr(response, "candidates", []) or []
        for candidate in candidates:
            content = getattr(candidate, "content", None)
            parts = getattr(content, "parts", []) if content else []
            joined = "".join(
                getattr(part, "text", "") or "" for part in parts
            )
            if joined.strip():
                return joined
            finish = getattr(candidate, "finish_reason", None)
            if finish:
                logger.warning("candidate finish_reason: %s", finish)
    except Exception as fallback_err:
        logger.warning("candidate extraction failed: %s", fallback_err)

    return ""


def generate_ai_text(prompt):

    if not GEMINI_API_KEY:

        return (
            "Backend Error: GEMINI_API_KEY is not set "
            "in environment variables."
        )

    _ensure_gemini_configured()

    last_error = None

    for model_name in GEMINI_MODEL_CHAIN:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            text = _extract_response_text(response)
            if text.strip():
                logger.info("Gemini reply via %s (%d chars)", model_name, len(text))
                return text
            last_error = Exception(f"empty response from {model_name}")
            logger.warning("empty response from %s, trying next model", model_name)
            continue
        except Exception as e:
            last_error = e
            logger.warning("%s failed: %s", model_name, e)
            continue

    final_err = str(last_error) if last_error else "all models unavailable"
    logger.error("All Gemini models failed: %s", final_err)
    return f"Gemini API Error: {final_err}"

def load_pdf_context():

    if not os.path.exists(PDF_CONTEXT_FILE):

        return ""

    with open(
        PDF_CONTEXT_FILE,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:

        return file.read()[-5000:]


def fallback_chat_reply(question: str) -> str:

    text = (question or "").strip()
    lower_text = text.lower()

    if lower_text in {"hi", "hello", "hey", "hii"}:

        return (
            "Hi! I am here. The AI service is unavailable right now, "
            "but the chat is working."
        )

    if (
        lower_text.startswith("who is ") or
        lower_text.startswith("what is ")
    ):

        return (
            "I cannot look that up properly because the AI service is "
            "unavailable right now. Please try again in a moment after "
            "checking the backend API key and server logs."
        )

    return (
        "I received your message, but the AI service is unavailable "
        "right now. Please try again in a moment."
    )

# =========================
# REQUEST MODELS
# =========================

class ChatRequest(BaseModel):
    message: str
    history: list = []

class AuthRequest(BaseModel):
    email: str
    password: str
class CodeRequest(BaseModel):
    code: str  
    user_input: str = ""  
class EmailRequest(BaseModel):

    senderEmail: str

    senderPassword: str

    recipientEmail: str

    description: str = ""

    length: str = "medium"

class DocumentRequest(BaseModel):

    topic: str

    doc_type: str = "document"

    length: str = "medium"

class SpeakRequest(BaseModel):

    text: str

    voice: str = "Fritz-PlayAI"

# =========================
# ROOT
# =========================

@app.get("/")
def root():

    return {
        "message": "Backend Running"
    }

# =========================
# HEALTH
# =========================

@app.get("/health")
def health():

    return {
        "status": "ok",
        "provider": "gemini"
    }

# =========================
# SIGNUP
# =========================

@app.post("/signup")
def signup(req: AuthRequest):

    db = SessionLocal()

    existing_user = db.query(User).filter(
        User.email == req.email
    ).first()

    if existing_user:

        return {
            "message": "User already exists"
        }

    hashed_password = hash_password(
        req.password
    )

    user = User(
        email=req.email,
        password=hashed_password
    )

    db.add(user)

    db.commit()

    return {
        "message": "Signup successful"
    }

# =========================
# LOGIN
# =========================

@app.post("/login")
def login(req: AuthRequest):

    db = SessionLocal()

    user = db.query(User).filter(
        User.email == req.email
    ).first()

    if not user:

        return {
            "message": "Invalid email"
        }

    valid = verify_password(
        req.password,
        user.password
    )

    if not valid:

        return {
            "message": "Invalid password"
        }

    token = create_token(user.email)

    return {
        "token": token
    }

# =========================
# CHAT
# =========================

@app.post("/chat")
async def chat(req: ChatRequest):

    try:

        question = req.message

        history = req.history or []

        # =========================
        # CONVERSATION MEMORY
        # =========================

        conversation_text = ""

        for msg in history[-6:]:

            sender = msg.get("sender", "")
            text = msg.get("text", "")

            if sender == "user":

                conversation_text += (
                    f"User: {text}\n"
                )

            else:

                conversation_text += (
                    f"Assistant: {text}\n"
                )

        # =========================
        # LIGHTWEIGHT PDF CONTEXT
        # =========================

        context = load_pdf_context()

        # =========================
        # FINAL PROMPT
        # =========================

        final_prompt = f"""
You are a helpful AI assistant.

IMPORTANT RULES:

1. If the user asks for code:
- ALWAYS give proper code
- ALWAYS use markdown code blocks
- ALWAYS specify language

Example:

```python
print("hello")
``` 
2. Use conversation history to remember user context.
3. Use PDF context ONLY if relevant.

Conversation History:
{conversation_text}

PDF Context:
{context}

Current User Question:
{question}

Answer naturally.
"""

        # =========================
        # GROQ RESPONSE
        # =========================

        reply = await asyncio.wait_for(
            asyncio.to_thread(
                generate_ai_text,
                final_prompt
            ),
            timeout=45
        )

        if not reply or not reply.strip():

            logger.error("Empty reply from generate_ai_text")
            return {
                "reply": fallback_chat_reply(question)
            }

        if reply.startswith("Backend Error:"):

            logger.error("Backend error reply: %s", reply)
            return {"reply": reply}

        if reply.startswith("Gemini API Error:"):

            logger.error("Gemini error reply: %s", reply)
            short = reply[:200]
            return {
                "reply": (
                    f"The AI service returned an error. {short}. "
                    "Please try again or check the Gemini API key/quota."
                )
            }

        return {
            "reply": reply
        }

    except asyncio.TimeoutError:

        logger.error("Gemini call timed out after 45s")
        return {
            "reply": (
                "The AI took too long to respond (over 45 seconds). "
                "Please try again."
            )
        }

    except Exception as e:

        logger.exception("Chat endpoint crashed")
        return {
            "reply": f"Backend error: {str(e)[:200]}"
        }
    
# =========================
# RUN PYTHON
# =========================

@app.post("/run-python")
async def run_python(req: CodeRequest):

    try:

        # STORE ORIGINAL STDIN/STDOUT

        old_stdout = sys.stdout
        old_stdin = sys.stdin

        # CAPTURE OUTPUT

        redirected_output = io.StringIO()

        sys.stdout = redirected_output

        # PASS USER INPUT

        sys.stdin = io.StringIO(
            req.user_input
        )

        # EXECUTE CODE

        exec(req.code)

        # GET OUTPUT

        output = redirected_output.getvalue()

        # RESTORE

        sys.stdout = old_stdout
        sys.stdin = old_stdin

        return {
            "output": output
        }

    except Exception as e:

        sys.stdout = old_stdout
        sys.stdin = old_stdin

        return {
            "output": str(e)
        }
    # =========================
# SEND EMAIL
# =========================

@app.post("/send-email")
async def send_email(
    req: EmailRequest
):

    try:

        # DEFAULT DESCRIPTION

        description = (
            req.description
            if req.description
            else
            "Professional email"
        )

        # DEFAULT LENGTH

        length = (
            req.length
            if req.length
            else
            "medium"
        )

        # =========================
        # AI EMAIL GENERATION
        # =========================

        email_prompt = f"""
Write a professional email.

Email topic:
{description}

Email length:
{length}

Generate:
1. Subject
2. Email body

Format:

Subject: ...

Body: ...
"""

        content = generate_ai_text(
            email_prompt
        )

        # =========================
        # SPLIT SUBJECT/BODY
        # =========================

        subject = "AI Generated Email"

        body = content

        if "Subject:" in content:

            parts = content.split(
                "Body:"
            )

            subject_part = (
                parts[0]
                .replace(
                    "Subject:",
                    ""
                )
                .strip()
            )

            body_part = (
                parts[1].strip()
                if len(parts) > 1
                else content
            )

            subject = subject_part

            body = body_part

        # =========================
        # EMAIL MESSAGE
        # =========================

        msg = MIMEMultipart()

        msg["From"] = (
            req.senderEmail
        )

        msg["To"] = (
            req.recipientEmail
        )

        msg["Subject"] = subject

        msg.attach(

            MIMEText(
                body,
                "plain"
            )
        )

        # =========================
        # SMTP GMAIL
        # =========================

        server = smtplib.SMTP(
            "smtp.gmail.com",
            587
        )

        server.starttls()

        server.login(
            req.senderEmail,
            req.senderPassword
        )

        server.send_message(msg)

        server.quit()

        return {

            "message":
            "Email sent successfully"
        }

    except Exception as e:

        return {

            "message":
            str(e)
        }
# =========================
# DOCUMENT GENERATION (LangChain tool wrapping Gemini)
# =========================

@tool
def generate_document_content(
    topic: str,
    doc_type: str,
    length: str,
) -> str:
    """Generate a professional document body via a Gemini-backed LangChain chain.

    Returns plain text: first line is the title, then a blank line,
    then the body paragraphs.
    """

    length_guide = {
        "short":  "Approximately 200 words.",
        "medium": "Approximately 500 words.",
        "long":   "Approximately 1000 words.",
    }.get(length, "Approximately 500 words.")

    llm = ChatGroq(
        model="llama-3.3-70b-versatile",
        groq_api_key=GROQ_API_KEY,
        temperature=0.5,
        timeout=45,
        max_retries=1,
    )

    prompt = ChatPromptTemplate.from_messages([
        (
            "system",
            "You are a professional writer producing clean, "
            "well-structured documents in a confident tone."
        ),
        (
            "human",
            "Write a {doc_type} about:\n\n{topic}\n\n"
            "Length: {length_guide}\n\n"
            "Format requirements:\n"
            "- The FIRST LINE must be a clear, descriptive title only "
            "(no quotes, no 'Title:' prefix).\n"
            "- Then a single blank line.\n"
            "- Then the body in proper paragraphs separated by blank lines.\n"
            "- Where appropriate, use bullet points starting with '- ' or "
            "numbered items like '1. ' — each on its own line.\n"
            "- Professional tone. NO markdown formatting "
            "(no **bold**, no ##headings, no ```code```).\n"
            "- Do not include any preamble like 'Here is your document'. "
            "Output ONLY the document."
        ),
    ])

    chain = prompt | llm | StrOutputParser()

    return chain.invoke({
        "topic": topic,
        "doc_type": doc_type,
        "length_guide": length_guide,
    })


def _safe_filename(name: str) -> str:

    cleaned = "".join(
        c if (c.isalnum() or c in (" ", "-", "_")) else ""
        for c in name
    ).strip()

    return (cleaned[:60] or "document")


def _generate_document_text(
    topic: str,
    doc_type: str,
    length: str,
) -> str:

    length_guide = {
        "short":  "Approximately 200 words.",
        "medium": "Approximately 500 words.",
        "long":   "Approximately 1000 words.",
    }.get(length, "Approximately 500 words.")

    prompt = f"""
Write a {doc_type} about:

{topic}

Length: {length_guide}

Format requirements:
- The FIRST LINE must be a clear, descriptive title only.
- Then a single blank line.
- Then the body in proper paragraphs separated by blank lines.
- Where appropriate, use bullet points starting with '- ' or
  numbered items like '1. ' each on its own line.
- Professional tone. No markdown formatting.
- Do not include any preamble like 'Here is your document'.
Output ONLY the document.
"""

    if GEMINI_API_KEY:

        return generate_ai_text(prompt)

    if GROQ_API_KEY:

        llm = ChatGroq(
            model="llama-3.3-70b-versatile",
            groq_api_key=GROQ_API_KEY,
            temperature=0.5,
            timeout=45,
            max_retries=1,
        )

        chat_prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                "You are a professional writer producing clean, "
                "well-structured documents in a confident tone."
            ),
            ("human", "{prompt}"),
        ])

        chain = chat_prompt | llm | StrOutputParser()

        return chain.invoke({
            "prompt": prompt,
        })

    raise ValueError(
        "Set GEMINI_API_KEY or GROQ_API_KEY in environment variables."
    )


def _parse_document_content(content: str) -> tuple[str, list[str]]:

    lines = [
        line.rstrip()
        for line in content.strip().split("\n")
        if line.strip()
    ]

    title = lines[0] if lines else "Document"
    body_lines = lines[1:] if len(lines) > 1 else []

    return title, body_lines


def _pdf_text(value: str) -> str:

    text = (
        value
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )

    return text.encode(
        "latin-1",
        errors="replace"
    ).decode("latin-1")


def _build_pdf(title: str, body_lines: list[str]) -> bytes:

    page_width = 612
    page_height = 792
    margin = 72
    max_width_chars = 86
    body_size = 11
    title_size = 20
    line_gap = 16

    pages = []
    current_page = []
    y = page_height - margin

    def add_page():

        nonlocal current_page, y

        if current_page:

            pages.append(current_page)

        current_page = []
        y = page_height - margin

    def add_line(
        text: str,
        size: int = body_size,
        font: str = "F1",
        indent: int = 0,
    ):

        nonlocal y

        if y < margin:

            add_page()

        current_page.append({
            "text": text,
            "size": size,
            "font": font,
            "x": margin + indent,
            "y": y,
        })

        y -= line_gap if size == body_size else 28

    for wrapped in textwrap.wrap(
        title,
        width=52
    ) or [title]:

        add_line(
            wrapped,
            size=title_size,
            font="F2"
        )

    y -= 10

    bullet_prefixes = ("- ", "* ", "• ")

    for raw_line in body_lines:

        text = raw_line.strip()

        if not text:

            y -= line_gap
            continue

        is_bullet = text.startswith(bullet_prefixes)
        line_text = f"- {text[2:].strip()}" if is_bullet else text
        indent = 12 if is_bullet else 0

        wrapped_lines = textwrap.wrap(
            line_text,
            width=max_width_chars - (4 if is_bullet else 0)
        ) or [line_text]

        for index, wrapped in enumerate(wrapped_lines):

            add_line(
                wrapped,
                indent=indent if index else 0
            )

        y -= 4

    add_page()

    if not pages:

        pages = [[{
            "text": title or "Document",
            "size": title_size,
            "font": "F2",
            "x": margin,
            "y": page_height - margin,
        }]]

    objects = []

    def add_object(payload: bytes) -> int:

        objects.append(payload)
        return len(objects)

    catalog_id = add_object(
        b"<< /Type /Catalog /Pages 2 0 R >>"
    )
    pages_id = add_object(b"")
    font_regular_id = add_object(
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    )
    font_bold_id = add_object(
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>"
    )

    page_ids = []

    for page_lines in pages:

        commands = ["BT"]

        for item in page_lines:

            commands.append(
                f"/{item['font']} {item['size']} Tf "
                f"1 0 0 1 {item['x']} {item['y']} Tm "
                f"({_pdf_text(item['text'])}) Tj"
            )

        commands.append("ET")

        stream = "\n".join(commands).encode(
            "latin-1",
            errors="replace"
        )

        stream_id = add_object(
            b"<< /Length " +
            str(len(stream)).encode("ascii") +
            b" >>\nstream\n" +
            stream +
            b"\nendstream"
        )

        page_id = add_object(
            (
                "<< /Type /Page /Parent 2 0 R "
                f"/MediaBox [0 0 {page_width} {page_height}] "
                f"/Contents {stream_id} 0 R "
                "/Resources << /Font << "
                f"/F1 {font_regular_id} 0 R "
                f"/F2 {font_bold_id} 0 R "
                ">> >> >>"
            ).encode("ascii")
        )

        page_ids.append(page_id)

    objects[pages_id - 1] = (
        "<< /Type /Pages /Kids [" +
        " ".join(
            f"{page_id} 0 R" for page_id in page_ids
        ) +
        f"] /Count {len(page_ids)} >>"
    ).encode("ascii")

    output = io.BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]

    for index, payload in enumerate(objects, start=1):

        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode("ascii"))
        output.write(payload)
        output.write(b"\nendobj\n")

    xref_at = output.tell()
    output.write(
        f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    )
    output.write(b"0000000000 65535 f \n")

    for offset in offsets[1:]:

        output.write(
            f"{offset:010d} 00000 n \n".encode("ascii")
        )

    output.write(
        (
            "trailer\n"
            f"<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n"
            "startxref\n"
            f"{xref_at}\n"
            "%%EOF\n"
        ).encode("ascii")
    )

    return output.getvalue()


@app.post("/generate-document")
async def generate_document(
    req: DocumentRequest
):

    try:

        if not GEMINI_API_KEY:

            return {
                "error":
                    "GEMINI_API_KEY is not set "
                    "in environment variables."
            }

        content = generate_document_content.invoke({
            "topic":    req.topic,
            "doc_type": req.doc_type or "document",
            "length":   req.length or "medium",
        })

        lines = [
            line.rstrip()
            for line in content.strip().split("\n")
            if line.strip()
        ]

        title = lines[0] if lines else "Document"
        body_lines = lines[1:] if len(lines) > 1 else []

        doc = Document()

        # 1-inch margins.
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # Body style.
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 1.4
        normal.paragraph_format.space_after  = Pt(8)

        # Heading 1 style.
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(22)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        bullet_prefixes  = ("- ", "* ", "• ")
        numbered_prefixes = tuple(
            f"{i}. " for i in range(1, 30)
        )

        for line in body_lines:

            text = line.strip()

            if text.startswith(bullet_prefixes):

                doc.add_paragraph(
                    text[2:].strip(),
                    style="List Bullet",
                )

            elif text.startswith(numbered_prefixes):

                _, _, after = text.partition(". ")

                doc.add_paragraph(
                    (after or text).strip(),
                    style="List Number",
                )

            else:

                doc.add_paragraph(text)

        buffer = io.BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        filename = f"{_safe_filename(title)}.docx"

        return StreamingResponse(
            buffer,
            media_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            headers={
                "Content-Disposition":
                    f'attachment; filename="{filename}"',
                "Access-Control-Expose-Headers":
                    "Content-Disposition",
            },
        )

    except Exception as e:

        return {
            "error": str(e)
        }


@app.post("/generate-pdf")
async def generate_pdf(
    req: DocumentRequest
):

    try:

        content = _generate_document_text(
            req.topic,
            req.doc_type or "document",
            req.length or "medium",
        )

        title, body_lines = _parse_document_content(
            content
        )

        pdf_bytes = _build_pdf(
            title,
            body_lines
        )

        filename = f"{_safe_filename(title)}.pdf"

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition":
                    f'attachment; filename="{filename}"',
                "Access-Control-Expose-Headers":
                    "Content-Disposition",
            },
        )

    except Exception as e:

        return {
            "error": str(e)
        }

# =========================
# TRANSCRIBE AUDIO (Groq Whisper)
# =========================

@app.post("/transcribe")
async def transcribe(
    file: UploadFile = File(...)
):

    try:

        client = get_groq_client()

        if client is None:

            return {
                "transcript": "",
                "error":
                    "GROQ_API_KEY is not set "
                    "in environment variables."
            }

        audio_bytes = await file.read()

        filename = (
            file.filename or "audio.webm"
        )

        content_type = (
            file.content_type or "audio/webm"
        )

        transcription = (
            client
            .audio
            .transcriptions
            .create(
                file=(
                    filename,
                    audio_bytes,
                    content_type,
                ),
                model="whisper-large-v3-turbo",
                response_format="json",
            )
        )

        transcript = (
            getattr(transcription, "text", "") or ""
        ).strip()

        return {
            "transcript": transcript
        }

    except Exception as e:

        return {
            "transcript": "",
            "error": str(e),
        }

# =========================
# TEXT-TO-SPEECH (Groq PlayAI TTS)
# =========================

@app.post("/speak")
async def speak(req: SpeakRequest):

    try:

        client = get_groq_client()

        if client is None:

            return {
                "error":
                    "GROQ_API_KEY is not set "
                    "in environment variables."
            }

        text = (req.text or "").strip()

        if not text:

            return {
                "error": "Empty text."
            }

        # Strip markdown noise that doesn't read well aloud.
        spoken = (
            text
            .replace("```", "")
            .replace("**", "")
            .replace("__", "")
            .replace("##", "")
            .replace("#", "")
        )

        # Groq's PlayAI TTS has a per-request character cap (~10k).
        spoken = spoken[:9500]

        voice = (
            req.voice or "Fritz-PlayAI"
        )

        response = (
            client
            .audio
            .speech
            .create(
                model="playai-tts",
                voice=voice,
                input=spoken,
                response_format="wav",
            )
        )

        audio_bytes = response.read()

        return StreamingResponse(
            io.BytesIO(audio_bytes),
            media_type="audio/wav",
            headers={
                "Cache-Control": "no-store",
            },
        )

    except Exception as e:

        return {
            "error": str(e),
        }

# =========================
# UPLOAD PDF
# =========================

@app.post("/upload-pdf")
async def upload_pdf(
    file: UploadFile = File(...)
):

    try:

        # CREATE PDF FOLDER

        os.makedirs(
            "pdfs",
            exist_ok=True
        )

        # SAVE FILE

        safe_pdf_name = os.path.basename(
            file.filename or "uploaded.pdf"
        )

        file_path = os.path.join(
            "pdfs",
            safe_pdf_name
        )

        with open(
            file_path,
            "wb"
        ) as buffer:

            shutil.copyfileobj(
                file.file,
                buffer
            )

        # READ PDF

        reader = PdfReader(
            file_path
        )

        text = ""

        for page in reader.pages:

            extracted = page.extract_text()

            if extracted:

                text += extracted

        # CHUNK TEXT

        chunks = [

            text[i:i+500]

            for i in range(
                0,
                len(text),
                500
            )
        ]

        # STORE LIGHTWEIGHT CONTEXT

        with open(
            PDF_CONTEXT_FILE,
            "a",
            encoding="utf-8"
        ) as context_file:

            context_file.write(
                "\n\n".join(chunks[:20])
            )

            context_file.write("\n\n")

        return {

            "message":
            "PDF uploaded successfully"
        }

    except Exception as e:

        return {

            "message":
            str(e)
        }
